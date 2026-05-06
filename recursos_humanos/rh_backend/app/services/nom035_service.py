# app/services/nom035_service.py
from datetime import datetime
from typing import List, Optional, Tuple
import json
from io import BytesIO

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from reportlab.lib import colors

from app.models.nom035 import (
    Nom035Cycle,
    Nom035Question,
    Nom035CycleQuestion,
    Nom035Submission,
    Nom035Answer,
)
from app.services.notification_service import notification_service





T_CYCLE = Nom035Cycle.__tablename__
T_QUESTION = Nom035Question.__tablename__
T_CYCLE_Q = Nom035CycleQuestion.__tablename__
T_SUB = Nom035Submission.__tablename__
T_ANS = Nom035Answer.__tablename__
T_USERS = "users"


def _dt_str(dt) -> Optional[str]:
    if not dt:
        return None
    if isinstance(dt, str):
        return dt
    try:
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return str(dt)


def parse_dt(s: Optional[str]) -> Optional[datetime]:
    if not s:
        return None
    ss = s.strip().replace("T", " ")
    try:
        return datetime.strptime(ss, "%Y-%m-%d %H:%M:%S")
    except ValueError:
        return datetime.fromisoformat(ss)


def require_admin(user: dict):
    role_id = int(user.get("role_id", 3))
    if role_id not in (1, 2):
        raise PermissionError("No autorizado (admin)")


def require_user(user: dict):
    user_id = int(user.get("id", 0) or 0)
    if user_id <= 0:
        raise PermissionError("No autorizado")


def _likert_options_1_to_5():
    return [
        {"id": 1, "option_text": "Nunca"},
        {"id": 2, "option_text": "Casi nunca"},
        {"id": 3, "option_text": "A veces"},
        {"id": 4, "option_text": "Casi siempre"},
        {"id": 5, "option_text": "Siempre"},
    ]


def _yes_no_options():
    return [
        {"id": 1, "option_text": "Sí"},
        {"id": 2, "option_text": "No"},
    ]


class Nom035Service:
    def _cursor(self, conn):
        return conn.cursor(dictionary=True)

    def _ensure_user_profile(self, cur, user_id: int):
        cur.execute(
            """
            INSERT INTO nom035_user_profile (user_id)
            VALUES (%s)
            ON DUPLICATE KEY UPDATE updated_at=CURRENT_TIMESTAMP
            """,
            (int(user_id),),
        )

    def _ensure_submission_profile(self, cur, submission_id: int, user_id: int):
        cur.execute(
            """
            INSERT INTO nom035_submission_profile (submission_id, user_id)
            VALUES (%s, %s)
            ON DUPLICATE KEY UPDATE user_id=VALUES(user_id)
            """,
            (int(submission_id), int(user_id)),
        )

    def _sync_profile_field(self, cur, user_id: int, submission_id: int, field: str, value):
        allowed_fields = {
            "sex",
            "age_range",
            "marital_status",
            "education_level",
            "job_type",
            "hiring_type",
            "staff_type",
            "workday_type",
            "shift_rotation",
            "time_current_position",
            "total_work_experience",
        }

        if field not in allowed_fields:
            raise ValueError(f"Campo no permitido: {field}")

        
        cur.execute(
            f"UPDATE nom035_user_profile SET {field}=%s WHERE user_id=%s",
            (value, int(user_id)),
        )


    def _build_text_report(self, detail: dict) -> str:
        user_info = detail["user"]
        cycle_info = detail["cycle"]
        sub_info = detail["submission"]
        sections = detail["sections"]

        def _safe(v):
            return "" if v is None else str(v).strip()

        def _ans(v):
            s = _safe(v)
            return s if s else "Sin contestar"

        lines = []
        lines.append("NOM-035 - GUIA DE REFERENCIA")
        lines.append(f"Ciclo: {_safe(cycle_info.get('title'))} ({_safe(cycle_info.get('year'))})")
        lines.append(
            f"Empleado: {_safe(user_info.get('name'))} | No: {_safe(user_info.get('employee_number'))}"
        )
        lines.append(
            f"Área: {_safe(user_info.get('area'))} | Puesto: {_safe(user_info.get('position'))}"
        )
        lines.append(f"Email: {_safe(user_info.get('email'))}")
        lines.append(
            f"Estatus: {_safe(sub_info.get('status'))} | Score: {_safe(sub_info.get('score_total'))} | Riesgo: {_safe(sub_info.get('risk_level'))}"
        )
        lines.append(f"Enviado: {_safe(sub_info.get('submitted_at'))}")
        lines.append("")

        for sec in sections:
            sec_title = _safe(sec.get("title")) or "GUÍA"
            lines.append(f"== {sec_title} ==")

            sec_ins = _safe(sec.get("instructions"))
            sec_desc = _safe(sec.get("description"))

            if sec_desc:
                lines.append(f"Descripción: {sec_desc}")
            if sec_ins:
                lines.append(f"Instrucciones: {sec_ins}")

            if sec.get("questions"):
                for q in sec.get("questions") or []:
                    meta = q.get("meta") or {}
                    lines.append(f"- {_safe(q.get('question_text'))}")
                    lines.append(f"  Respuesta: {_ans(q.get('answer_value'))}")

                    it = _safe(q.get("instruction_text") or meta.get("instruction_text"))
                    ht = _safe(q.get("help_text") or meta.get("help_text"))
                    dm = _safe(meta.get("dimension"))
                    dn = _safe(meta.get("domain"))
                    ct = _safe(meta.get("category"))

                    if ct:
                        lines.append(f"  Categoría: {ct}")
                    if dm:
                        lines.append(f"  Dimensión: {dm}")
                    if dn:
                        lines.append(f"  Dominio: {dn}")
                    if it:
                        lines.append(f"  Instrucción: {it}")
                    if ht:
                        lines.append(f"  Ayuda: {ht}")

                    lines.append("")

            elif sec.get("groups"):
                for grp in sec.get("groups") or []:
                    grp_title = _safe(grp.get("title")) or "Grupo"
                    lines.append(f"--- Grupo: {grp_title} ---")

                    grp_ins = _safe(grp.get("instructions"))
                    if grp_ins:
                        lines.append(f"Instrucciones grupo: {grp_ins}")

                    for q in grp.get("questions") or []:
                        meta = q.get("meta") or {}
                        lines.append(f"- {_safe(q.get('question_text'))}")
                        lines.append(f"  Respuesta: {_ans(q.get('answer_value'))}")

                        it = _safe(q.get("instruction_text") or meta.get("instruction_text"))
                        ht = _safe(q.get("help_text") or meta.get("help_text"))
                        dm = _safe(meta.get("dimension"))
                        dn = _safe(meta.get("domain"))

                        if dm:
                            lines.append(f"  Dimensión: {dm}")
                        if dn:
                            lines.append(f"  Dominio: {dn}")
                        if it:
                            lines.append(f"  Instrucción: {it}")
                        if ht:
                            lines.append(f"  Ayuda: {ht}")

                        lines.append("")

            lines.append("")

        return "\n".join(lines).strip()

    def _exp_safe(self, v):
        return "" if v is None else str(v).strip()

    def _exp_answer(self, v):
        s = self._exp_safe(v)
        return s if s else "Sin contestar"

    def _exp_user_area(self, user_info: dict) -> str:
        return (
            self._exp_safe(user_info.get("area_name")) or
            self._exp_safe(user_info.get("area")) or
            self._exp_safe(user_info.get("department")) or
            "Sin área"
        )


    def _exp_user_position(self, user_info: dict) -> str:
        return (
            self._exp_safe(user_info.get("position_name")) or
            self._exp_safe(user_info.get("position")) or
            self._exp_safe(user_info.get("job_title")) or
            "Sin puesto"
        )
    
    def _is_guide_v_autofill_question(self, q: dict) -> bool:
        qid = int(q.get("id") or 0)
        return qid in (188, 189, 194, 195)

    
    def _exp_group_questions_by_guide(self, sections):
        guide_order = ["I", "II", "III", "IV", "V"]
        grouped = {g: [] for g in guide_order}

        for sec in sections or []:
            questions = sec.get("questions") or []
            for q in questions:
                meta = q.get("meta") or {}
                guide = self._exp_safe(
                    q.get("guide") or meta.get("guide")
                ).upper().replace("GUÍA", "").replace("GUIA", "").strip()

                if guide == "1":
                    guide = "I"
                elif guide == "2":
                    guide = "II"
                elif guide == "3":
                    guide = "III"
                elif guide == "4":
                    guide = "IV"
                elif guide == "5":
                    guide = "V"

                if guide in grouped:
                    grouped[guide].append(q)

            for grp in sec.get("groups") or []:
                for q in grp.get("questions") or []:
                    meta = q.get("meta") or {}
                    guide = self._exp_safe(
                        q.get("guide") or meta.get("guide")
                    ).upper().replace("GUÍA", "").replace("GUIA", "").strip()

                    if guide == "1":
                        guide = "I"
                    elif guide == "2":
                        guide = "II"
                    elif guide == "3":
                        guide = "III"
                    elif guide == "4":
                        guide = "IV"
                    elif guide == "5":
                        guide = "V"

                    if guide in grouped:
                        grouped[guide].append(q)

        return grouped



    def _export_xlsx(self, submission_id: int, detail: dict) -> dict:
        user_info = detail["user"]
        cycle_info = detail["cycle"]
        sub_info = detail["submission"]
        sections = detail["sections"]

        grouped = self._exp_group_questions_by_guide(sections)
        area_name = self._exp_user_area(user_info)
        position_name = self._exp_user_position(user_info)

        wb = Workbook()
        ws = wb.active
        ws.title = "NOM035"

        dark_blue = "1F4E78"
        soft_blue = "D9EAF7"
        red_fill = "FDE9E7"
        gold_fill = "FFF2CC"

        thin = Side(style="thin", color="D9D9D9")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        title_font = Font(bold=True, size=16, color="FFFFFF")
        white_bold = Font(bold=True, color="FFFFFF")
        bold_font = Font(bold=True)
        question_font = Font(bold=True)
        answer_red_font = Font(bold=True, color="C00000")
        section_font = Font(bold=True, size=13, color="1F1F1F")

        row = 1

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        c = ws.cell(row=row, column=1, value="NOM-035 - GUÍA DE REFERENCIA")
        c.font = title_font
        c.fill = PatternFill(fill_type="solid", fgColor=dark_blue)
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border
        row += 2

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        c = ws.cell(row=row, column=1, value="RESULTADOS")
        c.font = section_font
        c.fill = PatternFill(fill_type="solid", fgColor=gold_fill)
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border
        row += 1

        general_data = [
            ("Ciclo", f"{self._exp_safe(cycle_info.get('title'))} ({self._exp_safe(cycle_info.get('year'))})"),
            ("Empleado", self._exp_safe(user_info.get("name"))),
            ("Número de empleado", self._exp_safe(user_info.get("employee_number"))),
            ("Área", area_name),
            ("Posición", position_name),
            ("Email", self._exp_safe(user_info.get("email"))),
            ("Estatus", self._exp_safe(sub_info.get("status"))),
            ("Score total", self._exp_safe(sub_info.get("score_total"))),
            ("Nivel de riesgo", self._exp_safe(sub_info.get("risk_level"))),
            ("Enviado", self._exp_safe(sub_info.get("submitted_at"))),
        ]

        for label, value in general_data:
            ws.cell(row=row, column=1, value=label).font = bold_font
            ws.cell(row=row, column=1).fill = PatternFill(fill_type="solid", fgColor=soft_blue)
            ws.cell(row=row, column=1).border = border
            ws.cell(row=row, column=2, value=value)
            ws.cell(row=row, column=2).border = border
            row += 1

        row += 1

        for guide in ["I", "II", "III", "IV", "V"]:
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            c = ws.cell(row=row, column=1, value=f"GUÍA DE REFERENCIA {guide}")
            c.font = section_font
            c.fill = PatternFill(fill_type="solid", fgColor=soft_blue)
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = border
            row += 1

            headers = ["#", "Pregunta", "Respuesta", "Categoría", "Dominio", "Dimensión"]
            for col, h in enumerate(headers, start=1):
                cell = ws.cell(row=row, column=col, value=h)
                cell.fill = PatternFill(fill_type="solid", fgColor=dark_blue)
                cell.font = white_bold
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = border
            row += 1

            questions = grouped.get(guide, []) or []
            visible_questions = [
                q for q in questions
                if not self._is_guide_v_autofill_question(q)
            ]

            if not visible_questions:
                for col in range(1, 7):
                    ws.cell(row=row, column=col).border = border
                msg = ws.cell(row=row, column=1, value="Sin preguntas registradas para esta guía.")
                msg.font = answer_red_font
                msg.alignment = Alignment(horizontal="left", vertical="center")
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
                row += 2
                continue

            for idx, q in enumerate(visible_questions, start=1):
                meta = q.get("meta") or {}
                answer = self._exp_answer(q.get("answer_value"))
                is_empty = answer == "Sin contestar"

                values = [
                    idx,
                    self._exp_safe(q.get("question_text")),
                    answer,
                    self._exp_safe(meta.get("category")),
                    self._exp_safe(meta.get("domain")),
                    self._exp_safe(meta.get("dimension")),
                ]

                for col, value in enumerate(values, start=1):
                    cell = ws.cell(row=row, column=col, value=value)
                    cell.border = border
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

                    if col == 2:
                        cell.font = question_font

                    if col == 3:
                        cell.font = answer_red_font
                        if is_empty:
                            cell.fill = PatternFill(fill_type="solid", fgColor=red_fill)

                row += 1

            row += 1

        widths = {
            "A": 8,
            "B": 70,
            "C": 28,
            "D": 24,
            "E": 24,
            "F": 22,
        }
        for col, width in widths.items():
            ws.column_dimensions[col].width = width

        for r in ws.iter_rows():
            for cell in r:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        output = BytesIO()
        wb.save(output)
        data = output.getvalue()

        return {
            "filename": f"nom035_submission_{submission_id}.xlsx",
            "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "bytes": data,
        }

    


    def _export_docx(self, submission_id: int, detail: dict) -> dict:
        user_info = detail["user"]
        cycle_info = detail["cycle"]
        sub_info = detail["submission"]
        sections = detail["sections"]

        grouped = self._exp_group_questions_by_guide(sections)
        area_name = self._exp_user_area(user_info)
        position_name = self._exp_user_position(user_info)

        doc = Document()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("NOM-035 - GUÍA DE REFERENCIA")
        r.bold = True
        r.font.size = Pt(16)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = subtitle.add_run("RESULTADOS")
        rr.bold = True
        rr.font.size = Pt(13)

        info_lines = [
            ("Ciclo", f"{self._exp_safe(cycle_info.get('title'))} ({self._exp_safe(cycle_info.get('year'))})"),
            ("Empleado", self._exp_safe(user_info.get("name"))),
            ("Número de empleado", self._exp_safe(user_info.get("employee_number"))),
            ("Área", area_name),
            ("Posición", position_name),
            ("Email", self._exp_safe(user_info.get("email"))),
            ("Estatus", self._exp_safe(sub_info.get("status"))),
            ("Score total", self._exp_safe(sub_info.get("score_total"))),
            ("Nivel de riesgo", self._exp_safe(sub_info.get("risk_level"))),
            ("Enviado", self._exp_safe(sub_info.get("submitted_at"))),
        ]

        for label, value in info_lines:
            p = doc.add_paragraph()
            a = p.add_run(f"{label}: ")
            a.bold = True
            p.add_run(value)

        for guide in ["I", "II", "III", "IV", "V"]:
            doc.add_paragraph("")
            h = doc.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            x = h.add_run(f"GUÍA DE REFERENCIA {guide}")
            x.bold = True
            x.font.size = Pt(14)

            questions = grouped.get(guide, []) or []
            visible_questions = [
                q for q in questions
                if not self._is_guide_v_autofill_question(q)
            ]

            if not visible_questions:
                p = doc.add_paragraph()
                rr = p.add_run("Sin preguntas registradas para esta guía.")
                rr.bold = True
                rr.font.color.rgb = RGBColor(192, 0, 0)
                continue

            for idx, q in enumerate(visible_questions, start=1):
                meta = q.get("meta") or {}
                answer = self._exp_answer(q.get("answer_value"))

                p = doc.add_paragraph()
                qrun = p.add_run(f"{idx}. {self._exp_safe(q.get('question_text'))}")
                qrun.bold = True

                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Pt(18)
                p2.add_run("Respuesta: ").bold = True
                r2 = p2.add_run(answer)
                r2.bold = True
                r2.font.color.rgb = RGBColor(192, 0, 0)

                cat = self._exp_safe(meta.get("category"))
                dom = self._exp_safe(meta.get("domain"))
                dim = self._exp_safe(meta.get("dimension"))

                if cat:
                    x = doc.add_paragraph()
                    x.paragraph_format.left_indent = Pt(18)
                    x.add_run("Categoría: ").bold = True
                    x.add_run(cat)

                if dom:
                    x = doc.add_paragraph()
                    x.paragraph_format.left_indent = Pt(18)
                    x.add_run("Dominio: ").bold = True
                    x.add_run(dom)

                if dim:
                    x = doc.add_paragraph()
                    x.paragraph_format.left_indent = Pt(18)
                    x.add_run("Dimensión: ").bold = True
                    x.add_run(dim)

        output = BytesIO()
        doc.save(output)
        data = output.getvalue()

        return {
            "filename": f"nom035_submission_{submission_id}.docx",
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "bytes": data,
        }



    def _export_pdf(self, submission_id: int, detail: dict) -> dict:
        user_info = detail["user"]
        cycle_info = detail["cycle"]
        sub_info = detail["submission"]
        sections = detail["sections"]

        grouped = self._exp_group_questions_by_guide(sections)
        area_name = self._exp_user_area(user_info)
        position_name = self._exp_user_position(user_info)

        output = BytesIO()
        c = canvas.Canvas(output, pagesize=letter)
        width, height = letter

        left = 40
        right = width - 40
        top = height - 40
        bottom = 40
        max_width = right - left
        y = top

        def ensure_space(h=20):
            nonlocal y
            if y < bottom + h:
                c.showPage()
                y = top

        def draw_wrapped(text, font_name="Helvetica", font_size=10, color=colors.black, indent=0, gap=14):
            nonlocal y
            ensure_space(30)
            c.setFont(font_name, font_size)
            c.setFillColor(color)
            lines = simpleSplit(text or "", font_name, font_size, max_width - indent)
            if not lines:
                lines = [""]
            for line in lines:
                ensure_space(gap)
                c.drawString(left + indent, y, line)
                y -= gap

        def draw_center(text, font_name="Helvetica-Bold", font_size=13, fill_color=None, text_color=colors.black):
            nonlocal y
            ensure_space(28)
            if fill_color:
                c.setFillColor(fill_color)
                c.rect(left, y - 6, max_width, 20, fill=1, stroke=0)
            c.setFillColor(text_color)
            c.setFont(font_name, font_size)
            c.drawCentredString(width / 2, y, text)
            y -= 24

        draw_center(
            "NOM-035 - GUÍA DE REFERENCIA",
            font_size=16,
            fill_color=colors.HexColor("#1F4E78"),
            text_color=colors.white,
        )
        y -= 4

        draw_center(
            "RESULTADOS",
            font_size=12,
            fill_color=colors.HexColor("#FFF2CC"),
            text_color=colors.black,
        )

        general_data = [
            ("Ciclo", f"{self._exp_safe(cycle_info.get('title'))} ({self._exp_safe(cycle_info.get('year'))})"),
            ("Empleado", self._exp_safe(user_info.get("name"))),
            ("Número de empleado", self._exp_safe(user_info.get("employee_number"))),
            ("Área", area_name),
            ("Posición", position_name),
            ("Email", self._exp_safe(user_info.get("email"))),
            ("Estatus", self._exp_safe(sub_info.get("status"))),
            ("Score total", self._exp_safe(sub_info.get("score_total"))),
            ("Nivel de riesgo", self._exp_safe(sub_info.get("risk_level"))),
            ("Enviado", self._exp_safe(sub_info.get("submitted_at"))),
        ]

        for label, value in general_data:
            ensure_space(18)
            c.setFont("Helvetica-Bold", 10)
            c.setFillColor(colors.black)
            c.drawString(left, y, f"{label}:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 110, y, value)
            y -= 16

        y -= 8

        for guide in ["I", "II", "III", "IV", "V"]:
            draw_center(
                f"GUÍA DE REFERENCIA {guide}",
                font_size=12,
                fill_color=colors.HexColor("#D9EAF7"),
                text_color=colors.black,
            )

            questions = grouped.get(guide, []) or []
            visible_questions = [
                q for q in questions
                if not self._is_guide_v_autofill_question(q)
            ]

            if not visible_questions:
                draw_wrapped(
                    "Sin preguntas registradas para esta guía.",
                    font_name="Helvetica-Bold",
                    color=colors.red,
                )
                y -= 8
                continue

            for idx, q in enumerate(visible_questions, start=1):
                meta = q.get("meta") or {}
                answer = self._exp_answer(q.get("answer_value"))

                draw_wrapped(
                    f"{idx}. {self._exp_safe(q.get('question_text'))}",
                    font_name="Helvetica-Bold",
                    font_size=10,
                )
                draw_wrapped(
                    f"Respuesta: {answer}",
                    font_name="Helvetica-Bold",
                    font_size=10,
                    color=colors.red,
                    indent=14,
                )

                cat = self._exp_safe(meta.get("category"))
                dom = self._exp_safe(meta.get("domain"))
                dim = self._exp_safe(meta.get("dimension"))

                if cat:
                    draw_wrapped(f"Categoría: {cat}", indent=14)
                if dom:
                    draw_wrapped(f"Dominio: {dom}", indent=14)
                if dim:
                    draw_wrapped(f"Dimensión: {dim}", indent=14)

                y -= 6

        c.save()
        data = output.getvalue()

        return {
            "filename": f"nom035_submission_{submission_id}.pdf",
            "content_type": "application/pdf",
            "bytes": data,
        }



    def _answer_to_numeric(self, response_type, answer_value):
        av = str(answer_value or "").strip().upper()

        # Normaliza formatos como 1B, 0A, 4C
        if av and av[0].isdigit():
            av_num = av[0]
        else:
            av_num = av

        if response_type == "yes_no":
            # En tu BD real: 1 = Sí, 0 = No
            if av_num == "1":
                return 1
            if av_num == "0":
                return 0

            if av_num in ("SI", "SÍ", "YES", "TRUE"):
                return 1
            if av_num in ("NO", "FALSE"):
                return 0

            return 0

        if response_type == "likert":
            try:
                v = int(av_num)
            except Exception:
                return 0

            # Si viene 0..4
            if 0 <= v <= 4:
                return v

            # Si por alguna razón viene 1..5
            if 1 <= v <= 5:
                return max(0, min(4, v - 1))

            return 0

        return 0

    def _risk_level_from_score(self, score: int) -> str:
        if score >= 90:
            return "Muy Alto"
        if score >= 70:
            return "Alto"
        if score >= 45:
            return "Medio"
        if score >= 20:
            return "Bajo"
        return "Nulo"

    def _calculate_domain_scores(self, conn, submission_id: int):
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT
                    q.id AS question_id,
                    q.response_type,
                    q.domain,
                    COALESCE(q.reverse_scoring, 0) AS reverse_scoring,
                    COALESCE(q.weight, 1) AS weight,
                    a.answer_value
                FROM {T_ANS} a
                INNER JOIN {T_QUESTION} q ON q.id = a.question_id
                WHERE a.submission_id = %s
                  AND COALESCE(q.domain, '') <> ''
                  AND q.guide IN ('II', 'III')
                """,
                (int(submission_id),),
            )
            rows = cur.fetchall() or []

            domain_scores = {}

            for r in rows:
                domain = str(r.get("domain") or "").strip()
                if not domain:
                    continue

                response_type = str(r.get("response_type") or "").strip().lower()
                value = self._answer_to_numeric(response_type, r.get("answer_value"))

                reverse_scoring = int(r.get("reverse_scoring") or 0)
                if reverse_scoring == 1:
                    if response_type == "likert":
                        value = 4 - value
                    elif response_type == "yes_no":
                        value = 1 - value

                try:
                    weight = float(r.get("weight") or 1)
                except Exception:
                    weight = 1.0

                score = value * weight
                domain_scores.setdefault(domain, 0.0)
                domain_scores[domain] += score

            result = []
            for domain, score in domain_scores.items():
                score_int = int(round(score))
                result.append({
                    "domain": domain,
                    "score": score_int,
                    "risk_level": self._risk_level_from_score(score_int),
                })

            result.sort(key=lambda x: x["domain"].lower())
            return result

        finally:
            cur.close()

    def _save_domain_scores(self, conn, submission_id: int, domain_scores: list[dict]):
        cur = self._cursor(conn)
        try:
            cur.execute(
                """
                DELETE FROM nom035_submission_domain_scores
                WHERE submission_id = %s
                """,
                (int(submission_id),),
            )

            for item in domain_scores:
                cur.execute(
                    """
                    INSERT INTO nom035_submission_domain_scores
                        (submission_id, domain, score, risk_level)
                    VALUES (%s, %s, %s, %s)
                    """,
                    (
                        int(submission_id),
                        str(item.get("domain") or "").strip(),
                        int(item.get("score") or 0),
                        str(item.get("risk_level") or "Nulo"),
                    ),
                )
        finally:
            cur.close()

    # ---------------------------
    # USER
    # ---------------------------
    def list_active_cycles(self, conn) -> List[dict]:
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT id, year, title, start_at, due_at, status
                FROM {T_CYCLE}
                WHERE status = 'active'
                ORDER BY year DESC, id DESC
                """
            )
            return cur.fetchall()
        finally:
            cur.close()


    def list_visible_cycles_for_completed(self, conn) -> List[dict]:
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT id, year, title, start_at, due_at, status
                FROM {T_CYCLE}
                WHERE status IN ('active', 'closed')
                ORDER BY year DESC, id DESC
                """
            )
            return cur.fetchall()
        finally:
            cur.close()

    def ensure_submission(self, conn, cycle_id: int, user_id: int) -> dict:
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT *
                FROM {T_SUB}
                WHERE cycle_id=%s AND user_id=%s
                LIMIT 1
                """,
                (cycle_id, user_id),
            )
            sub = cur.fetchone()
            if sub:
                return sub

            cur.execute(
                f"""
                INSERT INTO {T_SUB} (cycle_id, user_id, status, locked)
                VALUES (%s, %s, 'available', 0)
                """,
                (cycle_id, user_id),
            )
            conn.commit()

            sub_id = cur.lastrowid
            cur.execute(f"SELECT * FROM {T_SUB} WHERE id=%s", (sub_id,))
            return cur.fetchone()
        finally:
            cur.close()

    def available_forms_for_user(self, conn, user_id: int) -> List[dict]:
        cycles = self.list_active_cycles(conn)
        out: List[dict] = []
        for c in cycles:
            sub = self.ensure_submission(conn, c["id"], user_id)

            if sub.get("status") == "submitted" or int(sub.get("locked") or 0) == 1:
                continue

            out.append(
                {
                    "form_id": c["id"],
                    "title": c["title"],
                    "type": "nom035",
                    "status": sub.get("status"),
                    "started_at": _dt_str(sub.get("started_at")),
                    "due_at": _dt_str(c.get("due_at")),
                    "submission_id": sub["id"],
                    "scope": "general",
                    "area": "General",
                }
            )
        return out

    def completed_forms_for_user(self, conn, user_id: int) -> List[dict]:
        cycles = self.list_visible_cycles_for_completed(conn)
        out: List[dict] = []

        for c in cycles:
            sub = self.ensure_submission(conn, c["id"], user_id)

            if sub.get("status") == "submitted" or int(sub.get("locked") or 0) == 1:
                out.append(
                    {
                        "form_id": c["id"],
                        "title": c["title"],
                        "type": "nom035",
                        "status": "submitted",
                        "submitted_at": _dt_str(sub.get("submitted_at")),
                        "score": sub.get("score_total"),
                        "observations": sub.get("observations") or "",
                        "submission_id": sub["id"],
                    }
                )

        return out

    def start_form(self, conn, cycle_id: int, user_id: int) -> dict:
        cur = self._cursor(conn)
        cur.execute(f"SELECT * FROM {T_CYCLE} WHERE id=%s LIMIT 1", (cycle_id,))
        cycle = cur.fetchone()
        if not cycle or cycle.get("status") != "active":
            cur.close()
            raise ValueError("Ciclo no disponible")

        sub = self.ensure_submission(conn, cycle_id, user_id)
        if int(sub.get("locked") or 0) == 1 or sub.get("status") == "submitted":
            cur.close()
            return {"submission_id": sub["id"], "status": "submitted"}

        if sub.get("status") != "in_progress":
            started_at = sub.get("started_at") or datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
            cur.execute(
                f"""
                UPDATE {T_SUB}
                SET status='in_progress', started_at=%s
                WHERE id=%s
                """,
                (started_at, sub["id"]),
            )
            conn.commit()
            sub["status"] = "in_progress"

        try:
            cur.execute(
                f"""
                SELECT id, area, position
                FROM {T_USERS}
                WHERE id=%s
                LIMIT 1
                """,
                (int(user_id),),
            )
            u = cur.fetchone() or {}

            raw_area = u.get("area")
            raw_position = u.get("position")

            area_id = None
            position_id = None

            if raw_area is not None and str(raw_area).strip() != "":
                raw_area_s = str(raw_area).strip()
                ar = None

                if raw_area_s.isdigit():
                    cur.execute(
                        "SELECT id, name FROM areas WHERE id=%s AND active=1 LIMIT 1",
                        (int(raw_area_s),),
                    )
                    ar = cur.fetchone()
                else:
                    cur.execute(
                        "SELECT id, name FROM areas WHERE name=%s AND active=1 LIMIT 1",
                        (raw_area_s,),
                    )
                    ar = cur.fetchone()

                if ar:
                    area_id = int(ar["id"])

            if raw_position is not None and str(raw_position).strip() != "":
                raw_pos_s = str(raw_position).strip()
                pr = None

                if raw_pos_s.isdigit():
                    cur.execute(
                        "SELECT id, name FROM positions WHERE id=%s AND active=1 LIMIT 1",
                        (int(raw_pos_s),),
                    )
                    pr = cur.fetchone()
                else:
                    if area_id is not None:
                        cur.execute(
                            """
                            SELECT id, name
                            FROM positions
                            WHERE name=%s AND area_id=%s AND active=1
                            LIMIT 1
                            """,
                            (raw_pos_s, int(area_id)),
                        )
                    else:
                        cur.execute(
                            "SELECT id, name FROM positions WHERE name=%s AND active=1 LIMIT 1",
                            (raw_pos_s,),
                        )
                    pr = cur.fetchone()

                if pr:
                    position_id = int(pr["id"])

            default_questionnaire_no = "1"
            default_applied_date = datetime.utcnow().date().isoformat()

            self._ensure_user_profile(cur, int(user_id))
            self._ensure_submission_profile(cur, int(sub["id"]), int(user_id))

            # PERFIL GLOBAL DEL USUARIO:
            # aquí sí queremos reflejar el estado actual
            cur.execute(
                """
                INSERT INTO nom035_user_profile (
                    user_id, questionnaire_no, applied_date, area_id, position_id
                )
                VALUES (%s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                    questionnaire_no = VALUES(questionnaire_no),
                    applied_date     = VALUES(applied_date),
                    area_id          = VALUES(area_id),
                    position_id      = VALUES(position_id),
                    updated_at       = CURRENT_TIMESTAMP
                """,
                (
                    int(user_id),
                    default_questionnaire_no,
                    default_applied_date,
                    area_id,
                    position_id,
                ),
            )

            # SNAPSHOT DEL SUBMISSION:
            # se inserta una vez y ya no se reescribe
            cur.execute(
                """
                INSERT INTO nom035_submission_profile (
                    submission_id, user_id, questionnaire_no, applied_date, area_id, position_id
                )
                VALUES (%s, %s, %s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                    user_id = VALUES(user_id)
                """,
                (
                    int(sub["id"]),
                    int(user_id),
                    default_questionnaire_no,
                    default_applied_date,
                    area_id,
                    position_id,
                ),
            )

            conn.commit()

        except Exception as e:
            conn.rollback()
            print(f"ERROR start_form() profile init user_id={user_id}: {e}")

        finally:
            cur.close()

        return {"submission_id": sub["id"], "status": sub["status"]}

    def get_form_detail(self, conn, cycle_id: int, user_id: int) -> dict:
        cur = self._cursor(conn)

        cur.execute(f"SELECT * FROM {T_CYCLE} WHERE id=%s LIMIT 1", (int(cycle_id),))
        cycle = cur.fetchone()
        if not cycle:
            cur.close()
            raise ValueError("Ciclo no existe")

        sub = self.ensure_submission(conn, int(cycle_id), int(user_id))
        submission_id = int(sub["id"])

        cur.execute(
            f"""
            SELECT
                cq.order_no AS cq_order,
                q.id,
                q.question_text,
                q.response_type,
                q.options_json,
                q.guide,
                q.category,
                q.order_no
            FROM {T_CYCLE_Q} cq
            JOIN {T_QUESTION} q ON q.id = cq.question_id
            WHERE cq.cycle_id=%s AND q.is_active=1
            ORDER BY cq.order_no ASC, q.order_no ASC, q.id ASC
            """,
            (int(cycle_id),),
        )
        rows = cur.fetchall() or []

        cur.execute(
            f"""
            SELECT question_id, answer_value
            FROM {T_ANS}
            WHERE submission_id=%s
            """,
            (int(submission_id),),
        )
        arows = cur.fetchall() or []
        answers_map = {int(a["question_id"]): (a.get("answer_value") or "") for a in arows}

        # =========================================================
        # LEER PERFIL DEL SUBMISSION (SNAPSHOT DEL CUESTIONARIO)
        # NO DEL USER_PROFILE
        # =========================================================
        cur.execute(
            """
            SELECT
                sp.questionnaire_no,
                sp.applied_date,
                sp.area_id,
                sp.position_id,
                a.name AS area_name,
                p.name AS position_name
            FROM nom035_submission_profile sp
            LEFT JOIN areas a ON a.id = sp.area_id
            LEFT JOIN positions p ON p.id = sp.position_id
            WHERE sp.submission_id=%s
            LIMIT 1
            """,
            (int(submission_id),),
        )
        prof = cur.fetchone()

        # Si no existe el snapshot del submission, intenta construirlo desde users
        if not prof:
            cur.execute(
                "SELECT area, position FROM users WHERE id=%s LIMIT 1",
                (int(user_id),),
            )
            u = cur.fetchone() or {}

            raw_area = u.get("area")
            raw_position = u.get("position")

            area_id = None
            position_id = None
            area_name = None
            position_name = None

            if raw_area is not None and str(raw_area).strip() != "":
                raw_area_s = str(raw_area).strip()
                ar = None

                if raw_area_s.isdigit():
                    cur.execute(
                        "SELECT id, name FROM areas WHERE id=%s AND active=1 LIMIT 1",
                        (int(raw_area_s),),
                    )
                    ar = cur.fetchone()
                else:
                    cur.execute(
                        "SELECT id, name FROM areas WHERE name=%s AND active=1 LIMIT 1",
                        (raw_area_s,),
                    )
                    ar = cur.fetchone()

                if ar:
                    area_id = int(ar["id"])
                    area_name = (ar.get("name") or "").strip()

            if raw_position is not None and str(raw_position).strip() != "":
                raw_pos_s = str(raw_position).strip()
                pr = None

                if raw_pos_s.isdigit():
                    cur.execute(
                        "SELECT id, name FROM positions WHERE id=%s AND active=1 LIMIT 1",
                        (int(raw_pos_s),),
                    )
                    pr = cur.fetchone()
                else:
                    if area_id is not None:
                        cur.execute(
                            """
                            SELECT id, name
                            FROM positions
                            WHERE name=%s AND area_id=%s AND active=1
                            LIMIT 1
                            """,
                            (raw_pos_s, int(area_id)),
                        )
                    else:
                        cur.execute(
                            "SELECT id, name FROM positions WHERE name=%s AND active=1 LIMIT 1",
                            (raw_pos_s,),
                        )
                    pr = cur.fetchone()

                if pr:
                    position_id = int(pr["id"])
                    position_name = (pr.get("name") or "").strip()

            prof = {
                "questionnaire_no": "1",
                "applied_date": datetime.utcnow().date(),
                "area_id": area_id,
                "position_id": position_id,
                "area_name": area_name,
                "position_name": position_name,
            }

        area_id = prof.get("area_id")
        position_id = prof.get("position_id")
        area_name = (prof.get("area_name") or "").strip()
        position_name = (prof.get("position_name") or "").strip()

        questionnaire_no = str(prof.get("questionnaire_no") or "1")

        raw_applied_date = prof.get("applied_date")
        if hasattr(raw_applied_date, "isoformat"):
            applied_date = raw_applied_date.isoformat()
        else:
            applied_date = str(raw_applied_date or "")

        user_area_opt = (
            [{"id": int(area_id), "option_text": area_name}]
            if area_id is not None and area_name
            else []
        )
        user_pos_opt = (
            [{"id": int(position_id), "option_text": position_name}]
            if position_id is not None and position_name
            else []
        )

        # =========================================================
        # AQUÍ YA NO USAMOS default_applied_date = hoy
        # NI LEEMOS DESDE user_profile
        # =========================================================
        answers_map[188] = questionnaire_no
        answers_map[189] = applied_date

        if position_id is not None:
            answers_map[194] = str(position_id)

        if area_id is not None:
            answers_map[195] = str(area_id)

        guide_instructions = {
            "I": (
                "El contenido de esta guía es un complemento para la mejor comprensión de la presente Norma y no es de cumplimiento obligatorio. En esta guía, se presenta un ejemplo de cuestionario que permite identificar a los trabajadores que han sido sujetos a acontecimientos traumáticos severos y que requieren valoración clínica.\n "
                "CUESTIONARIO PARA IDENTIFICAR A LOS TRABAJADORES QUE FUERON SUJETOS A ACONTECIMIENTOS TRAUMÁTICOS SEVEROS\n"
                "Seleccione la opción que corresponda a su respuesta en cada preguntas."
            ),
            "II": (
                "IDENTIFICACIÓN Y ANÁLISIS DE LOS FACTORES DE RIESGO PSICOSOCIAL\n"
                "El contenido de esta guía es un complemento para la mejor comprensión de la presente Norma, "
                "y no es de cumplimiento obligatorio,\n"
                "puede ser utilizada por aquellos centros de trabajo que tengan hasta 50 trabajadores.\n"
                "Seleccione la opción que corresponda a su respuesta en cada pregunta.\n"
            ),
            "III": (
                "IDENTIFICACIÓN Y ANÁLISIS DE LOS FACTORES DE RIESGO PSICOSOCIAL Y EVALUACIÓN DEL ENTORNO ORGANIZACIONAL\n"
                "El contenido de esta guía es un complemento para la mejor comprensión de la presente Norma, y no es de cumplimiento obligatorio.\n"
                "Seleccione la opción que corresponda a su respuesta en cada pregunta.\n"
            ),
            "IV": (
                "POLÍTICA DE PREVENCIÓN DE RIESGOS PSICOSOCIALES\n"
                "El contenido de esta guía es un complemento para la mejor comprensión de la presente Norma, y no es de cumplimiento obligatorio.\n"
            ),
            "V": "DATOS DEL TRABAJADOR\n",
        }

        category_titles = {
            "I": "I.- Acontecimiento traumático severo",
            "II": "II.- Recuerdos persistentes sobre el acontecimiento (durante el último mes):",
            "III": "III.- Esfuerzo por evitar circunstancias parecidas o asociadas al acontecimiento (durante el último mes):",
            "IV": "IV Política De Prevención De Riesgos Psicosociales:",
            "V": "V Datos del Trabajador:",
        }

        def _norm_guide(raw) -> str:
            s = (raw or "").strip().upper()
            if not s:
                return "I"
            s = s.replace("GUÍA", "").replace("GUIA", "").strip()
            s = s.replace("_", "").replace("-", "").strip()
            if s.startswith("G"):
                s = s[1:].strip()
            num_map = {"1": "I", "2": "II", "3": "III", "4": "IV", "5": "V"}
            if s in num_map:
                return num_map[s]
            if s in ("I", "II", "III", "IV", "V"):
                return s
            return "I"

        def _norm_cat_for_guide(guide: str, cat_raw: str) -> str:
            cat = (cat_raw or "").strip()
            if not cat:
                return "Sin apartado"
            if guide == "I":
                c = cat.strip().upper()
                return category_titles.get(c, cat)
            return cat

        questions_out = []

        for r in rows:
            rt = (r.get("response_type") or "").strip().lower()

            q_type = "text"
            options = []

            if rt == "likert":
                q_type = "single"
                options = _likert_options_1_to_5()

            elif rt == "yes_no":
                q_type = "single"
                options = _yes_no_options()

            elif rt == "multiple":
                q_type = "multi"

                raw = r.get("options_json") or []
                if isinstance(raw, str):
                    try:
                        raw = json.loads(raw)
                    except Exception:
                        raw = []

                opts = []
                if isinstance(raw, list):
                    for idx, item in enumerate(raw, start=1):
                        if isinstance(item, dict):
                            oid = item.get("id", item.get("value", idx))
                            txt = item.get("option_text", item.get("label", item.get("text", "")))
                            txt = str(txt).strip()
                            if txt:
                                opts.append({"id": oid, "option_text": txt})
                        else:
                            txt = str(item).strip()
                            if txt:
                                opts.append({"id": idx, "option_text": txt})
                options = opts

            elif rt == "open":
                q_type = "text"
                options = []

            qid = int(r["id"])
            guide = (r.get("guide") or "").strip().upper()

            if guide == "V":
                if qid == 188:
                    q_type = "single"
                    options = [{"id": questionnaire_no, "option_text": questionnaire_no}]
                elif qid == 189:
                    q_type = "single"
                    options = [{"id": applied_date, "option_text": applied_date}]
                elif qid == 194:
                    q_type = "single"
                    options = user_pos_opt
                elif qid == 195:
                    q_type = "single"
                    options = user_area_opt

            questions_out.append(
                {
                    "id": qid,
                    "question_text": r.get("question_text"),
                    "question_type": q_type,
                    "required": 1,
                    "options": options,
                    "answer_value": answers_map.get(qid, ""),
                    "meta": {
                        "guide": r.get("guide"),
                        "category": r.get("category"),
                        "order_no": r.get("cq_order"),
                    },
                }
            )

        sections_map = {}
        for q in questions_out:
            guide_raw = (q.get("meta") or {}).get("guide")
            cat_raw = (q.get("meta") or {}).get("category")

            g = _norm_guide(guide_raw)
            cat = _norm_cat_for_guide(g, cat_raw)

            if g not in sections_map:
                sections_map[g] = {
                    "title": f"GUÍA DE REFERENCIA {g}",
                    "instructions": guide_instructions.get(g),
                    "groups": {},
                }

            sections_map[g]["groups"].setdefault(cat, []).append(q)

        for gfix in ["I", "II", "III", "IV", "V"]:
            sections_map.setdefault(
                gfix,
                {
                    "title": f"GUÍA DE REFERENCIA {gfix}",
                    "instructions": guide_instructions.get(gfix),
                    "groups": {},
                },
            )

        sections_map["V"]["groups"].setdefault("DATOS DEL TRABAJADOR", [])
        sections_map["I"]["groups"].setdefault("GR.I (Reglas de aplicación)", [])

        policy_text_iv = (
            "En este centro de trabajo (Razón Social) en relación con la prevención de los factores de riesgo psicosocial; "
            "la prevención de la violencia laboral, y la promoción de un entorno organizacional favorable, se asumen los compromisos siguientes:\n\n\n"
            "• Es obligación de supervisores, gerentes y directores aplicar esta política y predicar con el ejemplo;\n\n"
            "• Los actos de violencia laboral no son tolerados, así como ningún incidente que propicie factores de riesgo psicosocial o acciones en contra del entorno organizacional favorable\n\n"
            "• Se aplican medidas encaminadas a la prevención de los factores de riesgo psicosocial; la prevención de la violencia laboral, y la promoción de un entorno organizacional favorable, para prevenir sus consecuencias adversas;\n\n"
            "• Se cuenta con un procedimiento de atención justo, que no permite represalias y evita reclamaciones abusivas o carentes de fundamento, y que garantiza la confidencialidad de los casos; \n\n"
            "• Se realizan acciones de sensibilización, programas de información y capacitación;\n\n"
            "• Se divulgan de forma eficaz las políticas de prevención y las medidas de prevención;\n\n"
            "• Todos los trabajadores participan para establecer y poner en práctica esta política en el lugar de trabajo;\n\n"
            "• Se respeta al ejercicio de los derechos del personal para observar sus creencias o prácticas o para satisfacer sus necesidades relacionadas con la raza, sexo, religión, etnia o edad o cualquier otra condición que pueda dar origen a la discriminación, y \n\n"
            "• Se crean espacios de participación y consulta, teniendo en cuenta las ideas de los trabajadores y empleados.\n\n\n"
            "• La política deberá indicar a los responsables (con capacidad para su aplicación) y los recursos que se disponen para cumplir con las políticas en todos los planos de la organización; la asignación de responsabilidades tanto a las personas como a los equipos de trabajo, los cuales reciben la capacitación para la aplicación de las políticas."
        )

        sections_map["IV"]["instructions"] = (
            ((sections_map["IV"].get("instructions") or "").rstrip() + "\n\n" + policy_text_iv).strip()
        )

        IV_GROUP_TITLE = "POLÍTICA DE PREVENCIÓN DE RIESGOS PSICOSOCIALES"
        sections_map["IV"]["groups"].setdefault(IV_GROUP_TITLE, [])

        group_instructions_map = {
            ("IV", IV_GROUP_TITLE): policy_text_iv,
            ("I", "GR.I (Reglas de aplicación)"): "",
        }

        guide_order = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5}
        groups_out = []

        for g, sec in sorted(sections_map.items(), key=lambda x: guide_order.get(x[0], 99)):
            cats_out = []

            def _cat_sort_key(kv):
                cat_name, qs = kv
                first_ord = (qs[0].get("meta", {}).get("order_no", 9999) if qs else 999999)
                return (first_ord, cat_name)

            for cat_name, qs in sorted(sec["groups"].items(), key=_cat_sort_key):
                cats_out.append(
                    {
                        "title": cat_name,
                        "instructions": group_instructions_map.get((g, cat_name)),
                        "questions": qs,
                    }
                )

            groups_out.append(
                {
                    "title": sec["title"],
                    "instructions": sec["instructions"],
                    "groups": cats_out,
                }
            )

        cur.close()
        return {"form": cycle, "sections": groups_out, "submission_id": submission_id}
  
   
    def upsert_answer(self, conn, submission_id: int, user_id: int, question_id: int, answer_value) -> dict:
        cur = self._cursor(conn)

        try:
            cur.execute(f"SELECT * FROM {T_SUB} WHERE id=%s LIMIT 1", (int(submission_id),))
            sub = cur.fetchone()
            if not sub:
                raise ValueError("Submission no existe")
            if int(sub.get("user_id") or 0) != int(user_id):
                raise PermissionError("No autorizado")
            if int(sub.get("locked") or 0) == 1 or sub.get("status") == "submitted":
                raise PermissionError("Cuestionario bloqueado")

            qid = int(question_id)
            val_str = "" if answer_value is None else str(answer_value).strip()

            # Asegura perfiles
            self._ensure_user_profile(cur, int(user_id))
            self._ensure_submission_profile(cur, int(submission_id), int(user_id))

            # -------------------------------------------------
            # GUÍA V: CAMPOS FIJOS / SOLO LECTURA
            # No deben sobrescribirse al guardar respuestas
            # -------------------------------------------------
            fixed_guide_v_questions = {188, 189, 194, 195}

            # Si quieres que ni siquiera se guarde respuesta en nom035_answers
            # para esas preguntas fijas, salimos aquí.
            if qid in fixed_guide_v_questions:
                conn.commit()
                return {"ok": True, "skipped": True}

            # -------------------------------------------------
            # Guardar o actualizar respuesta normal
            # -------------------------------------------------
            cur.execute(
                f"""
                SELECT id FROM {T_ANS}
                WHERE submission_id=%s AND question_id=%s
                LIMIT 1
                """,
                (int(submission_id), qid),
            )
            existing = cur.fetchone()

            if existing:
                cur.execute(
                    f"UPDATE {T_ANS} SET answer_value=%s WHERE id=%s",
                    (str(answer_value), int(existing["id"])),
                )
            else:
                cur.execute(
                    f"""
                    INSERT INTO {T_ANS} (submission_id, question_id, answer_value)
                    VALUES (%s, %s, %s)
                    """,
                    (int(submission_id), qid, str(answer_value)),
                )

            # -------------------------------------------------
            # Sincronizar SOLO campos editables de perfil
            # -------------------------------------------------
            if qid == 190:
                self._sync_profile_field(cur, user_id, submission_id, "sex", val_str or None)

            elif qid == 191:
                self._sync_profile_field(cur, user_id, submission_id, "age_range", val_str or None)

            elif qid == 192:
                self._sync_profile_field(cur, user_id, submission_id, "marital_status", val_str or None)

            elif qid == 193:
                self._sync_profile_field(cur, user_id, submission_id, "education_level", val_str or None)

            elif qid == 196:
                self._sync_profile_field(cur, user_id, submission_id, "job_type", val_str or None)

            elif qid == 197:
                self._sync_profile_field(cur, user_id, submission_id, "hiring_type", val_str or None)

            elif qid == 198:
                self._sync_profile_field(cur, user_id, submission_id, "staff_type", val_str or None)

            elif qid == 199:
                self._sync_profile_field(cur, user_id, submission_id, "workday_type", val_str or None)

            elif qid == 200:
                self._sync_profile_field(cur, user_id, submission_id, "shift_rotation", val_str or None)

            elif qid == 201:
                self._sync_profile_field(cur, user_id, submission_id, "time_current_position", val_str or None)

            elif qid == 202:
                self._sync_profile_field(cur, user_id, submission_id, "total_work_experience", val_str or None)

            conn.commit()
            return {"ok": True}

        except Exception:
            conn.rollback()
            raise
        finally:
            cur.close()

    def _simple_score_from_answers(self, conn, submission_id: int) -> Tuple[int, str]:
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT
                    a.answer_value,
                    q.response_type,
                    q.guide,
                    COALESCE(q.reverse_scoring, 0) AS reverse_scoring
                FROM {T_ANS} a
                JOIN {T_QUESTION} q ON q.id = a.question_id
                WHERE a.submission_id=%s
                AND q.guide IN ('II', 'III')
                """,
                (int(submission_id),),
            )
            rows = cur.fetchall() or []

            total = 0
            answered_count = 0

            for r in rows:
                raw_answer = str(r.get("answer_value") or "").strip()
                if raw_answer == "":
                    continue

                response_type = str(r.get("response_type") or "").strip().lower()
                value = self._answer_to_numeric(response_type, r.get("answer_value"))

                reverse_scoring = int(r.get("reverse_scoring") or 0)
                if reverse_scoring == 1:
                    if response_type == "likert":
                        value = 4 - value
                    elif response_type == "yes_no":
                        value = 1 - value

                total += value
                answered_count += 1

            if answered_count == 0:
                return 0, "Sin contestar"

            level = self._risk_level_from_score(total)
            return total, level
        finally:
            cur.close()

    def submit_form(self, conn, submission_id: int, user_id: int) -> dict:
        cur = self._cursor(conn)

        try:
            submission_id = int(submission_id)
            user_id = int(user_id)

            cur.execute(
                f"""
                SELECT *
                FROM {T_SUB}
                WHERE id=%s
                LIMIT 1
                """,
                (submission_id,),
            )
            sub = cur.fetchone()

            if not sub:
                raise ValueError("Submission no existe")

            if int(sub.get("user_id") or 0) != user_id:
                raise PermissionError("No autorizado")

            if int(sub.get("locked") or 0) == 1 or sub.get("status") == "submitted":
                return {
                    "submission_id": submission_id,
                    "status": "submitted",
                    "score_total": sub.get("score_total"),
                    "risk_level": sub.get("risk_level"),
                }

            score, level = self._simple_score_from_answers(conn, submission_id)
            domain_scores = self._calculate_domain_scores(conn, submission_id)
            category_scores = self._calculate_category_scores(conn, submission_id)

            submitted_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")

            score_db = None if level == "Sin contestar" else score
            risk_db = None if level == "Sin contestar" else level

            cur.execute(
                f"""
                UPDATE {T_SUB}
                SET
                    score_total=%s,
                    risk_level=%s,
                    status='submitted',
                    locked=1,
                    submitted_at=%s
                WHERE id=%s
                """,
                (score_db, risk_db, submitted_at, submission_id),
            )

            self._save_domain_scores(conn, submission_id, domain_scores)
            self._save_category_scores(conn, submission_id, category_scores)

            conn.commit()

            try:
                cur.execute("CALL sp_nom035_snapshot_profile(%s)", (submission_id,))
                conn.commit()
            except Exception as e:
                print(f"ERROR sp_nom035_snapshot_profile({submission_id}): {e}")

            return {
                "submission_id": submission_id,
                "status": "submitted",
                "score_total": score_db,
                "risk_level": risk_db,
                "domain_scores": domain_scores,
                "category_scores": category_scores,
            }

        except Exception:
            conn.rollback()
            raise
        finally:
            cur.close()


    def get_result(self, conn, submission_id: int, user: dict) -> dict:
        cur = self._cursor(conn)
        try:
            cur.execute(f"SELECT * FROM {T_SUB} WHERE id=%s LIMIT 1", (submission_id,))
            sub = cur.fetchone()
            if not sub:
                raise ValueError("Submission no existe")

            role_id = int(user.get("role_id", 3))
            user_id = int(user.get("id", 0))

            if role_id not in (1, 2) and int(sub.get("user_id") or 0) != user_id:
                raise PermissionError("No autorizado")

            return {
                "submission_id": sub["id"],
                "cycle_id": sub["cycle_id"],
                "user_id": sub["user_id"],
                "status": sub.get("status"),
                "score_total": sub.get("score_total"),
                "risk_level": sub.get("risk_level"),
                "observations": sub.get("observations"),
                "submitted_at": sub.get("submitted_at"),
            }
        finally:
            cur.close()

    # ---------------------------
    # ADMIN
    # ---------------------------
    def admin_create_cycle(self, conn, user: dict, payload: dict) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        try:
            year = int(payload.get("year"))
            title = str(payload.get("title") or "").strip()
            start_at = parse_dt(payload.get("start_at"))
            due_at = parse_dt(payload.get("due_at"))
            status = str(payload.get("status") or "draft").strip().lower()

            if not title:
                raise ValueError("El título es obligatorio")

            if start_at is None or due_at is None:
                raise ValueError("Las fechas de inicio y límite son obligatorias")

            if due_at < start_at:
                raise ValueError("La fecha límite no puede ser menor a la fecha de inicio")

            if status not in ("draft", "active", "closed"):
                status = "draft"

            # No duplicar año
            cur.execute(
                f"""
                SELECT id, title, status
                FROM {T_CYCLE}
                WHERE year = %s
                AND status <> 'deleted'
                LIMIT 1
                """,
                (year,),
            )
            same_year = cur.fetchone()
            if same_year:
                raise ValueError(
                    f"Ya existe un ciclo para el año {year}: "
                    f"{same_year.get('title')} ({same_year.get('status')})"
                )

            # Solo un ciclo activo
            if status == "active":
                cur.execute(
                    f"""
                    SELECT id, title, year
                    FROM {T_CYCLE}
                    WHERE status = 'active'
                    AND status <> 'deleted'
                    LIMIT 1
                    """
                )
                active_cycle = cur.fetchone()
                if active_cycle:
                    raise ValueError(
                        "Ya existe un ciclo activo: "
                        f"{active_cycle.get('title')} ({active_cycle.get('year')})"
                    )

            cur.execute(
                f"""
                INSERT INTO {T_CYCLE} (year, title, start_at, due_at, status)
                VALUES (%s, %s, %s, %s, %s)
                """,
                (year, title, start_at, due_at, status),
            )
            cycle_id = cur.lastrowid
            conn.commit()

            processed_count = 0

            # Si nace activo, notifica
            if status == "active":
                cur.execute(
                    """
                    SELECT id, name, email, phone
                    FROM users
                    WHERE active = 1
                    """
                )
                users = cur.fetchall() or []

                for u in users:
                    notify_user_id = int(u.get("id") or 0)
                    notify_name = (u.get("name") or "").strip()
                    notify_email = (u.get("email") or "").strip()
                    notify_phone = (u.get("phone") or "").strip()

                    try:
                        notification_service.notify_cycle_start(
                            conn=conn,
                            cycle_id=cycle_id,
                            user_id=notify_user_id,
                            name=notify_name,
                            email=notify_email,
                            phone=notify_phone,
                            cycle_title=f"NOM-035 STPS ({year})",
                            due_at=due_at,
                            force_resend=False,  # creación normal
                        )
                        processed_count += 1
                    except Exception as e:
                        print(f"ERROR notify_cycle_start user_id={notify_user_id}: {e}")

            return {
                "ok": True,
                "id": cycle_id,
                "status": status,
                "message": f"Ciclo creado correctamente. Notificaciones procesadas: {processed_count}"
            }

        except Exception:
            conn.rollback()
            raise
        finally:
            cur.close()
    
    
    def admin_update_cycle(self, conn, user: dict, cycle_id: int, payload: dict) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        try:
            cur.execute(
                f"""
                SELECT id, year, title, start_at, due_at, status
                FROM {T_CYCLE}
                WHERE id=%s
                """,
                (cycle_id,),
            )
            current_cycle = cur.fetchone()
            if not current_cycle:
                raise ValueError("Ciclo no encontrado")

            old_status = str(current_cycle.get("status") or "").strip().lower()

            new_year = int(payload.get("year", current_cycle.get("year")))
            new_title = str(payload.get("title", current_cycle.get("title")) or "").strip()
            new_start_at = (
                parse_dt(payload.get("start_at"))
                if "start_at" in payload
                else current_cycle.get("start_at")
            )
            new_due_at = (
                parse_dt(payload.get("due_at"))
                if "due_at" in payload
                else current_cycle.get("due_at")
            )
            new_status = str(
                payload.get("status", current_cycle.get("status")) or "draft"
            ).strip().lower()

            # opcional: forzar reenvío manual
            resend_notifications = bool(payload.get("resend_notifications", False))

            if not new_title:
                raise ValueError("El título es obligatorio")

            if new_start_at is None or new_due_at is None:
                raise ValueError("Las fechas de inicio y límite son obligatorias")

            if new_due_at < new_start_at:
                raise ValueError("La fecha límite no puede ser menor a la fecha de inicio")

            if new_status not in ("draft", "active", "closed"):
                new_status = "draft"

            # No duplicar año con otro ciclo
            cur.execute(
                f"""
                SELECT id, title, status
                FROM {T_CYCLE}
                WHERE year = %s
                AND id <> %s
                AND status <> 'deleted'
                LIMIT 1
                """,
                (new_year, cycle_id),
            )
            same_year = cur.fetchone()
            if same_year:
                raise ValueError(
                    f"Ya existe otro ciclo para el año {new_year}: "
                    f"{same_year.get('title')} ({same_year.get('status')})"
                )

            # Solo un activo
            if new_status == "active":
                cur.execute(
                    f"""
                    SELECT id, title, year
                    FROM {T_CYCLE}
                    WHERE status = 'active'
                    AND id <> %s
                    AND status <> 'deleted'
                    LIMIT 1
                    """,
                    (cycle_id,),
                )
                active_cycle = cur.fetchone()
                if active_cycle:
                    raise ValueError(
                        "Ya existe otro ciclo activo: "
                        f"{active_cycle.get('title')} ({active_cycle.get('year')})"
                    )

            cur.execute(
                f"""
                UPDATE {T_CYCLE}
                SET year=%s, title=%s, start_at=%s, due_at=%s, status=%s
                WHERE id=%s
                """,
                (new_year, new_title, new_start_at, new_due_at, new_status, cycle_id),
            )
            conn.commit()

            # Notificar si:
            # 1) pasa de draft/closed a active
            # 2) o si se fuerza reenvío manual
            should_notify = (
                (old_status != "active" and new_status == "active")
                or
                (new_status == "active" and resend_notifications)
            )

            processed_count = 0
            error_count = 0
            error_messages = []

            if should_notify:
                cur.execute(
                    """
                    SELECT id, name, email, phone
                    FROM users
                    WHERE active = 1
                    """
                )
                users = cur.fetchall() or []

                for u in users:
                    notify_user_id = int(u.get("id") or 0)
                    notify_name = (u.get("name") or "").strip()
                    notify_email = (u.get("email") or "").strip()
                    notify_phone = (u.get("phone") or "").strip()

                    try:
                        notification_service.notify_cycle_start(
                            conn=conn,
                            cycle_id=cycle_id,
                            user_id=notify_user_id,
                            name=notify_name,
                            email=notify_email,
                            phone=notify_phone,
                            cycle_title=f"NOM-035 STPS ({new_year})",
                            due_at=new_due_at,
                            force_resend=resend_notifications,
                        )
                        processed_count += 1

                    except Exception as e:
                        error_count += 1
                        msg = f"user_id={notify_user_id}: {e}"
                        error_messages.append(msg)
                        print(f"ERROR notify_cycle_start {msg}")

                if error_count > 0:
                    return {
                        "ok": True,
                        "message": (
                            f"Ciclo actualizado. Notificaciones procesadas: {processed_count}. "
                            f"Errores: {error_count}"
                        ),
                        "errors": error_messages[:20],
                    }

                return {
                    "ok": True,
                    "message": f"Ciclo actualizado. Notificaciones procesadas: {processed_count}"
                }

            return {
                "ok": True,
                "message": "Ciclo actualizado correctamente",
                "debug": {
                    "old_status": old_status,
                    "new_status": new_status,
                    "should_notify": should_notify,
                    "resend_notifications": resend_notifications,
                },
            }

        except Exception:
            conn.rollback()
            raise
        finally:
            cur.close()

    def admin_delete_cycle(self, conn, user: dict, cycle_id: int) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        try:
            # validar existencia
            cur.execute(f"SELECT id, status FROM {T_CYCLE} WHERE id=%s", (cycle_id,))
            row = cur.fetchone()

            if not row:
                raise Exception("Ciclo no encontrado")

            # opcional: evitar borrar activos
            if row.get("status") == "active":
                raise Exception("No puedes eliminar un ciclo activo. Ciérralo primero.")

           # 1. respuestas
            cur.execute(f"""
                DELETE a FROM {T_ANS} a
                JOIN {T_SUB} s ON s.id = a.submission_id
                WHERE s.cycle_id = %s
            """, (cycle_id,))

            # 2. submissions
            cur.execute(f"""
                DELETE FROM {T_SUB}
                WHERE cycle_id = %s
            """, (cycle_id,))

            # 3. relaciones ciclo-preguntas
            cur.execute(f"""
                DELETE FROM {T_CYCLE_Q}
                WHERE cycle_id = %s
            """, (cycle_id,))

            # 4. ciclo
            cur.execute(f"""
                DELETE FROM {T_CYCLE}
                WHERE id = %s
            """, (cycle_id,))

            conn.commit()

            return {"ok": True, "message": "Ciclo eliminado completamente"}

        except Exception:
            conn.rollback()
            raise

        finally:
            cur.close()

    def admin_list_cycles(self, conn, user: dict) -> List[dict]:
        require_admin(user)
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"SELECT * FROM {T_CYCLE} WHERE status <> %s ORDER BY year DESC, id DESC",
                ("deleted",)
            )
            rows = cur.fetchall()
            return [
                {
                    "id": r["id"],
                    "year": r["year"],
                    "title": r["title"],
                    "start_at": _dt_str(r.get("start_at")),
                    "due_at": _dt_str(r.get("due_at")),
                    "status": r.get("status"),
                }
                for r in rows
            ]
        finally:
            cur.close()

    def admin_list_questions(self, conn, user: dict) -> List[dict]:
        require_admin(user)
        cur = self._cursor(conn)
        try:
            cur.execute(f"SELECT * FROM {T_QUESTION} ORDER BY id DESC")
            return cur.fetchall()
        finally:
            cur.close()

    def admin_upsert_question(self, conn, user: dict, payload: dict) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        try:
            options = payload.get("options_json")
            if isinstance(options, (list, dict)):
                options = json.dumps(options, ensure_ascii=False)

            qid = payload.get("id")

            if qid:
                cur.execute(
                    f"""
                    UPDATE {T_QUESTION}
                    SET guide=%s,
                        category=%s,
                        question_text=%s,
                        response_type=%s,
                        options_json=%s,
                        order_no=%s,
                        is_active=%s,
                        instruction_text=%s,
                        help_text=%s,
                        dimension=%s,
                        domain=%s,
                        reverse_scoring=%s,
                        weight=%s
                    WHERE id=%s
                    """,
                    (
                        payload["guide"],
                        payload.get("category"),
                        payload["question_text"],
                        payload["response_type"],
                        options,
                        int(payload.get("order_no", 1)),
                        int(payload.get("is_active", 1)),
                        payload.get("instruction_text"),
                        payload.get("help_text"),
                        payload.get("dimension"),
                        payload.get("domain"),
                        int(payload.get("reverse_scoring", 0)),
                        float(payload.get("weight", 1.0)),
                        int(qid),
                    ),
                )
                conn.commit()
                return {"id": int(qid)}

            cur.execute(
                f"""
                INSERT INTO {T_QUESTION}
                    (
                        guide,
                        category,
                        question_text,
                        response_type,
                        options_json,
                        order_no,
                        is_active,
                        instruction_text,
                        help_text,
                        dimension,
                        domain,
                        reverse_scoring,
                        weight
                    )
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    payload["guide"],
                    payload.get("category"),
                    payload["question_text"],
                    payload["response_type"],
                    options,
                    int(payload.get("order_no", 1)),
                    int(payload.get("is_active", 1)),
                    payload.get("instruction_text"),
                    payload.get("help_text"),
                    payload.get("dimension"),
                    payload.get("domain"),
                    int(payload.get("reverse_scoring", 0)),
                    float(payload.get("weight", 1.0)),
                ),
            )
            conn.commit()
            return {"id": cur.lastrowid}
        finally:
            cur.close()

    def admin_set_cycle_questions(self, conn, user: dict, cycle_id: int, question_ids: List[int]) -> dict:
        require_admin(user)

        if not question_ids:
            return {"ok": True, "count": 0, "skipped": True, "message": "Lista vacía: no se modificó el ciclo."}

        cur = self._cursor(conn)

        try:
            cur.execute(f"SELECT id FROM {T_QUESTION} WHERE guide='V' AND is_active=1 ORDER BY order_no ASC, id ASC")
            v_ids = [int(r["id"]) for r in (cur.fetchall() or [])]

            final_ids = []
            seen = set()
            for qid in (list(question_ids) + v_ids):
                qid = int(qid)
                if qid not in seen:
                    seen.add(qid)
                    final_ids.append(qid)

            cur.execute(f"DELETE FROM {T_CYCLE_Q} WHERE cycle_id=%s", (cycle_id,))
            order_no = 1
            for qid in final_ids:
                cur.execute(
                    f"INSERT INTO {T_CYCLE_Q} (cycle_id, question_id, order_no) VALUES (%s, %s, %s)",
                    (cycle_id, qid, order_no),
                )
                order_no += 1

            conn.commit()
            return {"ok": True, "count": len(final_ids)}
        finally:
            cur.close()

    def admin_get_submission_detail(self, conn, user: dict, submission_id: int) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        cur.execute(
            f"""
            SELECT
                s.*,
                c.title AS cycle_title,
                c.year AS cycle_year,
                c.start_at AS cycle_start_at,
                c.due_at AS cycle_due_at,
                u.name AS user_name,
                u.employee_number,
                u.area,
                u.position,
                u.email,
                a.name AS area_name,
                p.name AS position_name
            FROM {T_SUB} s
            JOIN {T_CYCLE} c ON c.id = s.cycle_id
            JOIN {T_USERS} u ON u.id = s.user_id
            LEFT JOIN areas a ON a.id = u.area
            LEFT JOIN positions p ON p.id = u.position
            WHERE s.id=%s
            LIMIT 1
            """,
            (int(submission_id),),
        )
        sub = cur.fetchone()
        if not sub:
            cur.close()
            raise ValueError("Submission no existe")

        cycle_id = int(sub["cycle_id"])

        cur.execute(
            f"""
            SELECT
                cq.order_no AS cq_order,
                q.id,
                q.question_text,
                q.response_type,
                q.options_json,
                q.guide,
                q.category,
                q.order_no,
                q.help_text,
                q.instruction_text,
                q.dimension,
                q.domain,
                q.reverse_scoring,
                q.weight
            FROM {T_CYCLE_Q} cq
            JOIN {T_QUESTION} q ON q.id = cq.question_id
            WHERE cq.cycle_id=%s AND q.is_active=1
            ORDER BY cq.order_no ASC, q.order_no ASC, q.id ASC
            """,
            (cycle_id,),
        )
        qrows = cur.fetchall() or []

        cur.execute(
            f"""
            SELECT question_id, answer_value
            FROM {T_ANS}
            WHERE submission_id=%s
            """,
            (int(submission_id),),
        )
        arows = cur.fetchall() or []
        answers_map = {int(a["question_id"]): (a.get("answer_value") or "") for a in arows}

        guide_instructions = {
            "I": "GUÍA DE REFERENCIA I",
            "II": "GUÍA DE REFERENCIA II",
            "III": "GUÍA DE REFERENCIA III",
            "IV": "GUÍA DE REFERENCIA IV",
            "V": "GUÍA DE REFERENCIA V",
        }

        def _norm_guide(raw) -> str:
            s = (raw or "").strip().upper()
            if not s:
                return "I"
            s = s.replace("GUÍA", "").replace("GUIA", "").strip()
            s = s.replace("_", "").replace("-", "").strip()
            if s.startswith("G"):
                s = s[1:].strip()
            num_map = {"1": "I", "2": "II", "3": "III", "4": "IV", "5": "V"}
            if s in num_map:
                return num_map[s]
            if s in ("I", "II", "III", "IV", "V"):
                return s
            return "I"

        questions_out = []
        for r in qrows:
            rt = (r.get("response_type") or "").strip().lower()

            if rt == "likert":
                q_type = "single"
                options = _likert_options_1_to_5()
            elif rt == "yes_no":
                q_type = "single"
                options = _yes_no_options()
            elif rt == "multiple":
                q_type = "multi"
                raw = r.get("options_json") or []
                if isinstance(raw, str):
                    try:
                        raw = json.loads(raw)
                    except Exception:
                        raw = []

                options = []
                if isinstance(raw, list):
                    for idx, item in enumerate(raw, start=1):
                        if isinstance(item, dict):
                            oid = item.get("id", item.get("value", idx))
                            txt = item.get("option_text", item.get("label", item.get("text", "")))
                            txt = str(txt).strip()
                            if txt:
                                options.append({"id": oid, "option_text": txt})
                        else:
                            txt = str(item).strip()
                            if txt:
                                options.append({"id": idx, "option_text": txt})
            else:
                q_type = "text"
                options = []

            qid = int(r["id"])
            raw_answer = (answers_map.get(qid) or "").strip()
            display_answer = self._display_answer_value(r, raw_answer)

            questions_out.append(
                {
                    "id": qid,
                    "question_text": r.get("question_text"),
                    "question_type": q_type,
                    "response_type": rt,
                    "required": 1,
                    "options": options,
                    "answer_value": display_answer,
                    "raw_answer_value": raw_answer,
                    "help_text": r.get("help_text"),
                    "instruction_text": r.get("instruction_text"),
                    "meta": {
                        "guide": r.get("guide"),
                        "category": r.get("category"),
                        "order_no": r.get("cq_order"),
                        "dimension": r.get("dimension"),
                        "domain": r.get("domain"),
                        "reverse_scoring": int(r.get("reverse_scoring") or 0),
                        "weight": float(r.get("weight") or 1.0),
                        "help_text": r.get("help_text"),
                        "instruction_text": r.get("instruction_text"),
                    },
                }
            )

        sections_map = {}
        for q in questions_out:
            meta = q.get("meta") or {}
            g = _norm_guide(meta.get("guide"))
            cat = str(meta.get("category") or "").strip() or "Sin apartado"

            if g not in sections_map:
                sections_map[g] = {
                    "title": f"GUÍA DE REFERENCIA {g}",
                    "instructions": guide_instructions.get(g, ""),
                    "groups": {},
                }

            sections_map[g]["groups"].setdefault(cat, []).append(q)

        for gfix in ["I", "II", "III", "IV", "V"]:
            sections_map.setdefault(
                gfix,
                {
                    "title": f"GUÍA DE REFERENCIA {gfix}",
                    "instructions": guide_instructions.get(gfix, ""),
                    "groups": {},
                },
            )

        guide_order = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5}
        sections = []

        for g, sec in sorted(sections_map.items(), key=lambda x: guide_order.get(x[0], 99)):
            groups_out = []

            def _cat_sort_key(kv):
                cat_name, qs = kv
                first_ord = (qs[0].get("meta", {}).get("order_no", 9999) if qs else 999999)
                return (first_ord, cat_name)

            for cat_name, qs in sorted(sec["groups"].items(), key=_cat_sort_key):
                groups_out.append(
                    {
                        "title": cat_name,
                        "instructions": None,
                        "questions": qs,
                    }
                )

            sections.append(
                {
                    "title": sec["title"],
                    "description": "",
                    "instructions": sec["instructions"],
                    "groups": groups_out,
                }
            )

        cur.close()
        return {
            "submission": {
                "id": int(sub["id"]),
                "cycle_id": int(sub["cycle_id"]),
                "status": sub.get("status"),
                "locked": int(sub.get("locked") or 0),
                "started_at": _dt_str(sub.get("started_at")),
                "submitted_at": _dt_str(sub.get("submitted_at")),
                "score_total": sub.get("score_total"),
                "risk_level": sub.get("risk_level"),
                "observations": sub.get("observations") or "",
            },
            "cycle": {
                "id": int(sub["cycle_id"]),
                "title": sub.get("cycle_title"),
                "year": sub.get("cycle_year"),
                "start_at": _dt_str(sub.get("cycle_start_at")),
                "due_at": _dt_str(sub.get("cycle_due_at")),
            },
            "user": {
                "id": int(sub.get("user_id") or 0),
                "name": sub.get("user_name"),
                "employee_number": sub.get("employee_number"),
                "area": sub.get("area"),
                "position": sub.get("position"),
                "area_name": sub.get("area_name"),
                "position_name": sub.get("position_name"),
                "email": sub.get("email"),
            },
            "sections": sections,
        }

    def admin_cycle_metrics(self, conn, user: dict, cycle_id: int) -> dict:
        require_admin(user)
        cur = self._cursor(conn)

        try:
            cur.execute(
                f"""
                SELECT status, COUNT(*) AS n
                FROM {T_SUB}
                WHERE cycle_id=%s
                GROUP BY status
                """,
                (int(cycle_id),),
            )
            by_status_rows = cur.fetchall() or []
            by_status = {r["status"]: int(r["n"]) for r in by_status_rows}

            cur.execute(
                f"""
                SELECT COALESCE(risk_level,'') AS risk_level, COUNT(*) AS n
                FROM {T_SUB}
                WHERE cycle_id=%s
                GROUP BY COALESCE(risk_level,'')
                """,
                (int(cycle_id),),
            )
            by_risk_rows = cur.fetchall() or []
            by_risk = {r["risk_level"] or "": int(r["n"]) for r in by_risk_rows}

            cur.execute(
                f"""
                SELECT
                  COUNT(*) AS n_scored,
                  AVG(score_total) AS avg_score
                FROM {T_SUB}
                WHERE cycle_id=%s AND score_total IS NOT NULL
                """,
                (int(cycle_id),),
            )
            avg_row = cur.fetchone() or {}
            n_scored = int(avg_row.get("n_scored") or 0)
            avg_score = avg_row.get("avg_score")
            avg_score = float(avg_score) if avg_score is not None else None

            cur.execute(
                f"SELECT COUNT(*) AS total FROM {T_SUB} WHERE cycle_id=%s",
                (int(cycle_id),),
            )
            total = int((cur.fetchone() or {}).get("total") or 0)

            submitted = int(by_status.get("submitted") or 0)
            in_progress = int(by_status.get("in_progress") or 0)
            available = int(by_status.get("available") or 0)

            completion_rate = (submitted / total) if total > 0 else 0.0

            return {
                "cycle_id": int(cycle_id),
                "totals": {
                    "total": total,
                    "submitted": submitted,
                    "in_progress": in_progress,
                    "available": available,
                    "completion_rate": completion_rate,
                },
                "by_status": by_status,
                "by_risk": by_risk,
                "scores": {
                    "n_scored": n_scored,
                    "avg_score": avg_score,
                },
            }
        finally:
            cur.close()

    def admin_export_submission(self, conn, user: dict, submission_id: int, fmt: str) -> dict:
        require_admin(user)

        detail = self.admin_get_submission_detail(conn, user, submission_id)
        fmt = (fmt or "").lower().strip()

        if fmt == "xlsx":
            return self._export_xlsx(submission_id, detail)

        if fmt == "docx":
            return self._export_docx(submission_id, detail)

        return self._export_pdf(submission_id, detail)

    def admin_list_sections(self, db, user: dict):
        role_id = int(user.get("role_id", 0) or 0)
        if role_id not in (1, 2):
            raise PermissionError("Sin permisos")

        cur = db.cursor(dictionary=True)
        try:
            cur.execute(
                """
                SELECT
                    id,
                    title,
                    description,
                    instructions,
                    order_no,
                    is_active,
                    created_at,
                    updated_at
                FROM nom035_sections
                WHERE is_active = 1
                ORDER BY order_no ASC, id ASC
                """
            )
            rows = cur.fetchall() or []
            return rows
        finally:
            cur.close()

    def admin_list_cycle_submissions(
        self,
        db,
        user: dict,
        cycle_id: int,
        status: str | None = None,
        q: str | None = None,
        risk: str | None = None,
        page: int = 1,
        page_size: int = 25,
    ):
        require_admin(user)

        cur = db.cursor(dictionary=True)

        try:
            page = max(int(page or 1), 1)
            page_size = max(min(int(page_size or 25), 200), 1)
            offset = (page - 1) * page_size

            filters = ["s.cycle_id = %s"]
            params = [int(cycle_id)]

            if status and str(status).strip():
                filters.append("s.status = %s")
                params.append(str(status).strip())

            if risk and str(risk).strip():
                filters.append("LOWER(COALESCE(s.risk_level, '')) = %s")
                params.append(str(risk).strip().lower())

            if q and str(q).strip():
                term = f"%{str(q).strip()}%"
                filters.append("""
                    (
                        COALESCE(u.name, '') LIKE %s
                        OR COALESCE(u.employee_number, '') LIKE %s
                        OR COALESCE(u.area, '') LIKE %s
                        OR COALESCE(u.position, '') LIKE %s
                    )
                """)
                params.extend([term, term, term, term])

            where_sql = " AND ".join(filters)

            cur.execute(f"""
                SELECT COUNT(*) AS total
                FROM nom035_submissions s
                INNER JOIN users u ON u.id = s.user_id
                WHERE {where_sql}
            """, tuple(params))
            total_row = cur.fetchone() or {}
            total = int(total_row.get("total") or 0)

            paginated_params = list(params)
            paginated_params.extend([page_size, offset])

            cur.execute(f"""
                SELECT
                    s.id AS submission_id,
                    s.user_id,
                    s.status,
                    s.score_total,
                    s.risk_level,
                    s.submitted_at,
                    u.id AS user_id_real,
                    u.name AS user_name,
                    u.employee_number,
                    u.area AS area_name,
                    u.position AS position_name
                FROM nom035_submissions s
                INNER JOIN users u ON u.id = s.user_id
                WHERE {where_sql}
                ORDER BY s.id DESC
                LIMIT %s OFFSET %s
            """, tuple(paginated_params))

            rows = cur.fetchall() or []
            submission_ids = [int(r["submission_id"]) for r in rows if r.get("submission_id")]

            category_map = {}
            domain_map = {}

            if submission_ids:
                placeholders = ",".join(["%s"] * len(submission_ids))

                try:
                    cur.execute(f"""
                        SELECT
                            submission_id,
                            domain,
                            score,
                            risk_level
                        FROM nom035_submission_domain_scores
                        WHERE submission_id IN ({placeholders})
                        ORDER BY submission_id, domain
                    """, tuple(submission_ids))
                    domain_rows = cur.fetchall() or []

                    for r in domain_rows:
                        sid = int(r["submission_id"])
                        domain_map.setdefault(sid, []).append({
                            "domain_name": r.get("domain"),
                            "score": float(r.get("score") or 0),
                            "risk_level": r.get("risk_level"),
                        })
                except Exception:
                    domain_map = {}

                try:
                    cur.execute(f"""
                        SELECT
                            submission_id,
                            category,
                            score,
                            risk_level
                        FROM nom035_submission_category_scores
                        WHERE submission_id IN ({placeholders})
                        ORDER BY submission_id, category
                    """, tuple(submission_ids))
                    category_rows = cur.fetchall() or []

                    for r in category_rows:
                        sid = int(r["submission_id"])
                        category_map.setdefault(sid, []).append({
                            "category_name": r.get("category"),
                            "score": float(r.get("score") or 0),
                            "risk_level": r.get("risk_level"),
                        })
                except Exception:
                    category_map = {}

            items = []
            for r in rows:
                sid = int(r["submission_id"])
                items.append({
                    "submission_id": sid,
                    "status": r.get("status"),
                    "score_total": float(r.get("score_total") or 0),
                    "risk_level": r.get("risk_level"),
                    "submitted_at": r.get("submitted_at").isoformat() if r.get("submitted_at") else None,
                    "user": {
                        "id": r.get("user_id_real"),
                        "name": r.get("user_name"),
                        "employee_number": r.get("employee_number"),
                        "area_id": None,
                        "position_id": None,
                        "area_name": r.get("area_name"),
                        "position_name": r.get("position_name"),
                    },
                    "category_results": category_map.get(sid, []),
                    "domain_results": domain_map.get(sid, []),
                })

            pending_users = []
            try:
                pending_sql = """
                    SELECT *
                    FROM (
                        SELECT
                            u.id,
                            u.name,
                            u.employee_number,
                            u.area AS area_name,
                            u.position AS position_name,
                            'not_seen' AS pending_type,
                            NULL AS status
                        FROM users u
                        LEFT JOIN nom035_submissions s2
                            ON s2.cycle_id = %s
                            AND s2.user_id = u.id
                        WHERE u.active = 1
                            AND COALESCE(u.role_id, 0) = 3
                            AND s2.id IS NULL

                        UNION ALL

                        SELECT
                            u.id,
                            u.name,
                            u.employee_number,
                            u.area AS area_name,
                            u.position AS position_name,
                            CASE
                                WHEN s2.status = 'available' THEN 'seen_not_answered'
                                WHEN s2.status = 'in_progress' THEN 'in_progress'
                                ELSE 'pending'
                            END AS pending_type,
                            s2.status
                        FROM users u
                        INNER JOIN nom035_submissions s2
                            ON s2.cycle_id = %s
                            AND s2.user_id = u.id
                        WHERE u.active = 1
                            AND COALESCE(u.role_id, 0) = 3
                            AND s2.status IN ('available', 'in_progress')
                    ) p
                """

                pending_params = [int(cycle_id), int(cycle_id)]

                pending_filters = []
                if q and str(q).strip():
                    term = f"%{str(q).strip()}%"
                    pending_filters.append("""
                        (
                            COALESCE(p.name, '') LIKE %s
                            OR COALESCE(p.employee_number, '') LIKE %s
                            OR COALESCE(p.area_name, '') LIKE %s
                            OR COALESCE(p.position_name, '') LIKE %s
                        )
                    """)
                    pending_params.extend([term, term, term, term])

                if pending_filters:
                    pending_sql += " WHERE " + " AND ".join(pending_filters)

                pending_sql += """
                    ORDER BY
                        CASE
                            WHEN p.pending_type = 'not_seen' THEN 1
                            WHEN p.pending_type = 'seen_not_answered' THEN 2
                            WHEN p.pending_type = 'in_progress' THEN 3
                            ELSE 4
                        END,
                        p.name ASC
                """

                cur.execute(pending_sql, tuple(pending_params))
                pending_rows = cur.fetchall() or []

                pending_users = [{
                    "id": r.get("id"),
                    "name": r.get("name"),
                    "employee_number": r.get("employee_number"),
                    "area_id": None,
                    "position_id": None,
                    "area_name": r.get("area_name"),
                    "position_name": r.get("position_name"),
                    "pending_type": r.get("pending_type"),
                    "status": r.get("status"),
                } for r in pending_rows]

            except Exception:
                pending_users = []

            cur.execute(f"""
                SELECT
                    s.id AS submission_id,
                    s.risk_level,
                    s.score_total,
                    s.status,
                    COALESCE(u.area, 'Sin área') AS area_name
                FROM nom035_submissions s
                INNER JOIN users u ON u.id = s.user_id
                WHERE {where_sql}
            """, tuple(params))
            summary_rows = cur.fetchall() or []

            summary_submission_ids = [int(r["submission_id"]) for r in summary_rows if r.get("submission_id")]

            summary_categories = []
            summary_domains = []

            if summary_submission_ids:
                placeholders = ",".join(["%s"] * len(summary_submission_ids))

                try:
                    cur.execute(f"""
                        SELECT
                            domain,
                            AVG(COALESCE(score, 0)) AS avg_score,
                            COUNT(*) AS total_rows
                        FROM nom035_submission_domain_scores
                        WHERE submission_id IN ({placeholders})
                        GROUP BY domain
                        ORDER BY avg_score DESC, domain ASC
                    """, tuple(summary_submission_ids))
                    summary_domains = [{
                        "name": r.get("domain"),
                        "avg_score": float(r.get("avg_score") or 0),
                        "count": int(r.get("total_rows") or 0),
                    } for r in (cur.fetchall() or [])]
                except Exception:
                    summary_domains = []

                try:
                    cur.execute(f"""
                        SELECT
                            category,
                            AVG(COALESCE(score, 0)) AS avg_score,
                            COUNT(*) AS total_rows
                        FROM nom035_submission_category_scores
                        WHERE submission_id IN ({placeholders})
                        GROUP BY category
                        ORDER BY avg_score DESC, category ASC
                    """, tuple(summary_submission_ids))
                    summary_categories = [{
                        "name": r.get("category"),
                        "avg_score": float(r.get("avg_score") or 0),
                        "count": int(r.get("total_rows") or 0),
                    } for r in (cur.fetchall() or [])]
                except Exception:
                    summary_categories = []

            risk_summary = {
                "muy_alto": 0,
                "alto": 0,
                "medio": 0,
                "bajo": 0,
                "nulo": 0,
            }

            departments_counter = {}
            heatmap = {}

            def norm_risk(label):
                s = str(label or "").strip().lower()
                if s == "muy alto":
                    return "muy_alto"
                if s == "alto":
                    return "alto"
                if s == "medio":
                    return "medio"
                if s == "bajo":
                    return "bajo"
                return "nulo"

            for r in summary_rows:
                rk = norm_risk(r.get("risk_level"))
                risk_summary[rk] += 1

                dept = r.get("area_name") or "Sin área"
                departments_counter[dept] = departments_counter.get(dept, 0) + 1

                if dept not in heatmap:
                    heatmap[dept] = {
                        "department": dept,
                        "muy_alto": 0,
                        "alto": 0,
                        "medio": 0,
                        "bajo": 0,
                        "nulo": 0,
                    }
                heatmap[dept][rk] += 1

            departments = [
                {"name": k, "count": v}
                for k, v in sorted(
                    departments_counter.items(),
                    key=lambda x: x[1],
                    reverse=True
                )
            ]

            heatmap_rows = sorted(
                heatmap.values(),
                key=lambda x: (
                    x["muy_alto"] + x["alto"] + x["medio"] + x["bajo"] + x["nulo"]
                ),
                reverse=True
            )

            compliance = {
                "policy_uploaded": False,
                "evidence_uploaded": False,
                "results_generated": len(summary_rows) > 0,
                "action_plan_created": False,
                "stps_file_ready": False,
            }

            return {
                "items": items,
                "pending_users": pending_users,
                "total": total,
                "page": page,
                "page_size": page_size,
                "summary": {
                    "risk": risk_summary,
                    "departments": departments,
                    "categories": summary_categories,
                    "domains": summary_domains,
                    "heatmap": heatmap_rows,
                    "compliance": compliance,
                },
            }

        finally:
            cur.close()

    def recalculate_submission_category_scores(self, db, submission_id: int):
        submission_id = int(submission_id)
        category_scores = self._calculate_category_scores(db, submission_id)
        self._save_category_scores(db, submission_id, category_scores)
        db.commit()
        return True

    def recalculate_submission_aggregates(self, db, submission_id: int):
        submission_id = int(submission_id)

        domain_scores = self._calculate_domain_scores(db, submission_id)
        category_scores = self._calculate_category_scores(db, submission_id)

        self._save_domain_scores(db, submission_id, domain_scores)
        self._save_category_scores(db, submission_id, category_scores)

        db.commit()
        return True

    def submit_nom035_submission(self, db, submission_id: int, user: dict):
        require_user(user)
        return self.submit_form(db, int(submission_id), int(user["id"]))

    def backfill_nom035_category_scores(self, db):
        cur = db.cursor(dictionary=True)
        try:
            cur.execute("""
                SELECT id
                FROM nom035_submissions
                WHERE status = 'submitted'
                ORDER BY id
            """)
            rows = cur.fetchall() or []

            total = 0
            for row in rows:
                submission_id = int(row["id"])
                self.recalculate_submission_aggregates(db, submission_id)
                total += 1

            db.commit()

            return {
                "ok": True,
                "processed": total,
            }

        except Exception:
            db.rollback()
            raise
        finally:
            cur.close()

    def _calculate_category_scores(self, conn, submission_id: int):
        cur = self._cursor(conn)
        try:
            cur.execute(
                f"""
                SELECT
                    q.id AS question_id,
                    q.response_type,
                    q.category,
                    COALESCE(q.reverse_scoring, 0) AS reverse_scoring,
                    COALESCE(q.weight, 1) AS weight,
                    a.answer_value
                FROM {T_ANS} a
                INNER JOIN {T_QUESTION} q ON q.id = a.question_id
                WHERE a.submission_id = %s
                    AND COALESCE(q.category, '') <> ''
                    AND q.guide IN ('II', 'III')
                """,
                (int(submission_id),),
            )
            rows = cur.fetchall() or []

            category_scores = {}

            for r in rows:
                category = str(r.get("category") or "").strip()
                if not category:
                    continue

                response_type = str(r.get("response_type") or "").strip().lower()
                value = self._answer_to_numeric(response_type, r.get("answer_value"))

                reverse_scoring = int(r.get("reverse_scoring") or 0)
                if reverse_scoring == 1:
                    if response_type == "likert":
                        value = 4 - value
                    elif response_type == "yes_no":
                        value = 1 - value

                try:
                    weight = float(r.get("weight") or 1)
                except Exception:
                    weight = 1.0

                score = value * weight
                category_scores.setdefault(category, 0.0)
                category_scores[category] += score

            result = []
            for category, score in category_scores.items():
                score_int = int(round(score))
                result.append({
                    "category": category,
                    "score": score_int,
                    "risk_level": self._risk_level_from_score(score_int),
                })

            result.sort(key=lambda x: x["category"].lower())
            return result

        finally:
            cur.close()

    def _save_category_scores(self, conn, submission_id: int, category_scores: list[dict]):
        cur = self._cursor(conn)
        try:
            cur.execute(
                """
                DELETE FROM nom035_submission_category_scores
                WHERE submission_id = %s
                """,
                (int(submission_id),),
            )

            for item in category_scores:
                cur.execute(
                    """
                    INSERT INTO nom035_submission_category_scores
                        (submission_id, category, score, risk_level)
                    VALUES (%s, %s, %s, %s)
                    """,
                    (
                        int(submission_id),
                        str(item.get("category") or "").strip(),
                        int(item.get("score") or 0),
                        str(item.get("risk_level") or "Nulo"),
                    ),
                )
        finally:
            cur.close()

    def _display_answer_value(self, question: dict, answer_value):
        raw = "" if answer_value is None else str(answer_value).strip()
        if raw == "":
            return ""

        rt = str(question.get("response_type") or "").strip().lower()

        if rt == "yes_no":
            v = raw.upper()
            if v in ("1", "SI", "SÍ", "YES", "TRUE"):
                return "Sí"
            if v in ("0", "2", "NO", "FALSE"):
                return "No"
            return raw

        if rt == "likert":
            v = raw.upper()

            mapping = {
                "0": "Nunca",
                "1": "Casi nunca",
                "2": "A veces",
                "3": "Casi siempre",
                "4": "Siempre",
                "5": "Siempre",
            }

            if v and v[0].isdigit():
                return mapping.get(v[0], raw)

            text_map = {
                "NUNCA": "Nunca",
                "CASI NUNCA": "Casi nunca",
                "A VECES": "A veces",
                "ALGUNAS VECES": "A veces",
                "CASI SIEMPRE": "Casi siempre",
                "SIEMPRE": "Siempre",
            }
            return text_map.get(v, raw)

        if rt == "multiple":
            raw_opts = question.get("options_json") or []
            if isinstance(raw_opts, str):
                try:
                    raw_opts = json.loads(raw_opts)
                except Exception:
                    raw_opts = []

            selected = []
            parsed_values = []

            try:
                if raw.startswith("[") and raw.endswith("]"):
                    arr = json.loads(raw)
                    if isinstance(arr, list):
                        parsed_values = [str(x).strip() for x in arr if str(x).strip()]
                else:
                    parsed_values = [x.strip() for x in raw.split(",") if x.strip()]
            except Exception:
                parsed_values = [raw]

            options = []
            if isinstance(raw_opts, list):
                for idx, item in enumerate(raw_opts, start=1):
                    if isinstance(item, dict):
                        oid = str(item.get("id", item.get("value", idx))).strip()
                        txt = str(
                            item.get("option_text", item.get("label", item.get("text", "")))
                        ).strip()
                        if txt:
                            options.append((oid, txt))
                    else:
                        txt = str(item).strip()
                        if txt:
                            options.append((str(idx), txt))

            for pv in parsed_values:
                found = None
                for oid, txt in options:
                    if pv == oid or pv.lower() == txt.lower():
                        found = txt
                        break
                selected.append(found or pv)

            return ", ".join(selected)

        return raw


nom035_service = Nom035Service()